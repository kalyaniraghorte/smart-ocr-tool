"""
Smart OCR Tool - Thonny Desktop App
==========================================
Supports : Images (PNG, JPG, BMP, TIFF, WEBP), PDF, DOCX, TXT
Features : OCR with preprocessing, Read aloud (Stop/Resume), Copy, Save

HOW RESUME WORKS:
  Text is split into sentences. Each sentence is spoken one at a time.
  When Stop is clicked, the current sentence index is saved.
  Resume picks up from that exact sentence — not from the beginning.

STEP 1 — Install Tesseract OCR engine (NOT a pip package):
  Windows : https://github.com/UB-Mannheim/tesseract/wiki
            Default path: C:/Program Files/Tesseract-OCR/tesseract.exe
  Mac     : brew install tesseract
  Linux   : sudo apt install tesseract-ocr

STEP 2 — Install Python libraries in Thonny:
  Tools > Manage Packages > install each:
    pytesseract
    Pillow
    opencv-python
    pdfplumber
    python-docx
    pyttsx3
"""

import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import threading
import os
import time
import re

# ── Library imports ───────────────────────────────────────────────────
try:
    import pytesseract
except ImportError:
    raise SystemExit("Missing: pytesseract\nThonny > Tools > Manage Packages > install 'pytesseract'")

try:
    from PIL import Image
except ImportError:
    raise SystemExit("Missing: Pillow\nThonny > Tools > Manage Packages > install 'Pillow'")

try:
    import cv2
    import numpy as np
except ImportError:
    raise SystemExit("Missing: opencv-python\nThonny > Tools > Manage Packages > install 'opencv-python'")

try:
    import pdfplumber
except ImportError:
    raise SystemExit("Missing: pdfplumber\nThonny > Tools > Manage Packages > install 'pdfplumber'")

try:
    import docx
except ImportError:
    raise SystemExit("Missing: python-docx\nThonny > Tools > Manage Packages > install 'python-docx'")

try:
    import pyttsx3
except ImportError:
    raise SystemExit("Missing: pyttsx3\nThonny > Tools > Manage Packages > install 'pyttsx3'")


# ── Tesseract path — Windows only ─────────────────────────────────────
if os.name == "nt":
    pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"


# ════════════════════════════════════════════════════════════════════
#  IMAGE PREPROCESSING
# ════════════════════════════════════════════════════════════════════
def preprocess_image(pil_image):
    w, h = pil_image.size
    pil_image = pil_image.resize((w * 2, h * 2), Image.LANCZOS)
    img_arr = np.array(pil_image.convert("RGB"))
    gray = cv2.cvtColor(img_arr, cv2.COLOR_RGB2GRAY)
    binary = cv2.adaptiveThreshold(
        gray, 255,
        cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
        cv2.THRESH_BINARY,
        blockSize=15, C=9
    )
    kernel = np.ones((1, 1), np.uint8)
    cleaned = cv2.morphologyEx(binary, cv2.MORPH_OPEN, kernel)
    return Image.fromarray(cleaned)


# ════════════════════════════════════════════════════════════════════
#  TEXT EXTRACTION
# ════════════════════════════════════════════════════════════════════
def extract_from_image(path):
    pil_img   = Image.open(path)
    processed = preprocess_image(pil_img)
    config    = "--oem 3 --psm 3"
    return pytesseract.image_to_string(processed, lang="eng", config=config).strip()


def extract_from_pdf(path, progress_cb=None):
    pages = []
    with pdfplumber.open(path) as pdf:
        total = len(pdf.pages)
        for i, page in enumerate(pdf.pages):
            if progress_cb:
                progress_cb(40 + int((i / total) * 45), f"Page {i+1} of {total}...")
            text = page.extract_text()
            if text and text.strip():
                pages.append(f"--- Page {i+1} ---\n{text.strip()}")
            else:
                img = page.to_image(resolution=250).original
                ocr = pytesseract.image_to_string(
                    preprocess_image(img), lang="eng", config="--oem 3 --psm 3")
                if ocr.strip():
                    pages.append(f"--- Page {i+1} (OCR) ---\n{ocr.strip()}")
    return "\n\n".join(pages)


def extract_from_docx(path):
    doc   = docx.Document(path)
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_txt = "\t".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_txt:
                lines.append(row_txt)
    return "\n".join(lines)


def extract_from_txt(path):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def extract_text(path, progress_cb=None):
    def step(p, m):
        if progress_cb:
            progress_cb(p, m)
        time.sleep(0.2)

    ext = os.path.splitext(path)[1].lower()
    step(8,  "Loading file...")

    if ext in (".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".webp"):
        step(20, "Upscaling image 2x...")
        step(38, "Applying adaptive threshold preprocessing...")
        step(52, "Running Tesseract OCR (English)...")
        text = extract_from_image(path)
    elif ext == ".pdf":
        step(20, "Opening PDF...")
        step(35, "Checking for text layer...")
        text = extract_from_pdf(path, progress_cb)
    elif ext in (".docx", ".doc"):
        step(25, "Reading Word document...")
        step(60, "Extracting paragraphs and tables...")
        text = extract_from_docx(path)
    elif ext == ".txt":
        step(45, "Reading plain text file...")
        text = extract_from_txt(path)
    else:
        raise ValueError(f"Unsupported file type: '{ext}'")

    step(95, "Finalizing...")
    step(100, "Done!")
    return text


# ════════════════════════════════════════════════════════════════════
#  SENTENCE SPLITTER
#  Splits text into individual sentences so we can track position
# ════════════════════════════════════════════════════════════════════
def split_into_sentences(text):
    """
    Split text into sentences using punctuation boundaries.
    Returns a list of non-empty sentence strings.
    """
    # Split on . ! ? followed by space or end of string
    raw = re.split(r'(?<=[.!?])\s+', text.strip())
    # Also split on newlines to handle bullet points / line-by-line content
    sentences = []
    for chunk in raw:
        for line in chunk.splitlines():
            line = line.strip()
            if line:
                sentences.append(line)
    return sentences if sentences else [text]


# ════════════════════════════════════════════════════════════════════
#  TTS ENGINE  — Sentence-by-sentence with true resume position
#
#  How it works:
#  - Text is split into a list of sentences
#  - A fresh pyttsx3 engine speaks ONE sentence at a time
#  - self._current_index tracks which sentence is being spoken
#  - stop() sets a stop flag, saving _current_index
#  - resume() calls speak_from(saved_index) to continue from there
#  - This gives true "resume from where you stopped" behaviour
# ════════════════════════════════════════════════════════════════════
class TTSEngine:
    def __init__(self):
        self.speaking       = False
        self._stop_flag     = threading.Event()
        self._thread        = None
        self._sentences     = []
        self._current_index = 0      # index of sentence currently being read

        # Get default speech rate from a temp engine
        _tmp = pyttsx3.init()
        self._base_rate = int(_tmp.getProperty("rate"))
        _tmp.stop()
        del _tmp

    def _speak_sentence(self, sentence):
        """
        Speak a single sentence using a fresh pyttsx3 engine.
        Returns True if completed normally, False if stopped mid-sentence.
        """
        completed = False
        try:
            engine = pyttsx3.init()
            engine.setProperty("rate", self._base_rate)
            engine.say(sentence)
            engine.startLoop(False)
            while not self._stop_flag.is_set():
                engine.iterate()
                if not engine.isBusy():
                    completed = True
                    break
            engine.endLoop()
            engine.stop()
            del engine
        except Exception:
            pass
        return completed

    def _run_from(self, start_index, on_done=None):
        """
        Internal thread: speak sentences starting from start_index.
        Saves position on stop so resume knows where to continue.
        """
        self.speaking = True
        self._current_index = start_index

        for i in range(start_index, len(self._sentences)):
            if self._stop_flag.is_set():
                # Stopped — save current position for resume
                self._current_index = i
                break
            self._current_index = i
            finished = self._speak_sentence(self._sentences[i])
            if not finished:
                # Stopped mid-sentence — stay on this sentence for resume
                self._current_index = i
                break
        else:
            # Loop completed all sentences naturally
            self._current_index = 0   # reset so next Read starts fresh
            self.speaking = False
            if on_done and not self._stop_flag.is_set():
                on_done()
            return

        self.speaking = False

    def speak(self, text, on_done=None):
        """Start reading from the beginning of the text."""
        self._sentences = split_into_sentences(text)
        self._current_index = 0
        self._start_thread(0, on_done)

    def resume(self, on_done=None):
        """Continue reading from the sentence where Stop was pressed."""
        if not self._sentences:
            return
        self._start_thread(self._current_index, on_done)

    def _start_thread(self, from_index, on_done=None):
        """Stop any existing thread, clear flag, start fresh thread."""
        self._stop_flag.set()
        if self._thread and self._thread.is_alive():
            self._thread.join(timeout=1.5)
        self._stop_flag.clear()
        self._thread = threading.Thread(
            target=self._run_from,
            args=(from_index, on_done),
            daemon=True
        )
        self._thread.start()

    def stop(self):
        """Stop speaking. Position is saved in _current_index."""
        self._stop_flag.set()
        self.speaking = False
        if self._thread and self._thread.is_alive():
            self._thread.join(timeout=1.5)
        self._thread = None

    def reset(self):
        """Fully reset — clears saved position."""
        self.stop()
        self._sentences     = []
        self._current_index = 0


# ════════════════════════════════════════════════════════════════════
#  THEME & FONTS
# ════════════════════════════════════════════════════════════════════
BG      = "#0f1117"
SURFACE = "#181b25"
SUR2    = "#1e2230"
BORDER  = "#2a2f42"
ACCENT  = "#6c63ff"
ACCENT2 = "#00d4aa"
TEXT    = "#e8eaf0"
MUTED   = "#8891a8"
DANGER  = "#ff5f6d"
WARNING = "#f59e0b"
WHITE   = "#ffffff"

F_HEAD  = ("Segoe UI", 11, "bold")
F_BODY  = ("Segoe UI", 10)
F_SMALL = ("Segoe UI", 9)
F_MONO  = ("Courier New", 10)
F_BTN   = ("Segoe UI", 10, "bold")


# ════════════════════════════════════════════════════════════════════
#  MAIN GUI
# ════════════════════════════════════════════════════════════════════
class App:
    def __init__(self, root):
        self.root      = root
        self.tts       = TTSEngine()
        self.filepath  = None

        self.root.title("Smart OCR Tool")
        self.root.geometry("920x780")
        self.root.minsize(700, 580)
        self.root.configure(bg=BG)

        style = ttk.Style()
        style.theme_use("clam")
        style.configure("TProgressbar",
                        troughcolor=BORDER, background=ACCENT,
                        bordercolor=BG, lightcolor=ACCENT, darkcolor=ACCENT)

        self._build_header()
        self._setup_scroll_area()
        self._build_upload()
        self._build_progress()
        self._build_output()

    # ── Header ────────────────────────────────────────────────────────
    def _build_header(self):
        hdr = tk.Frame(self.root, bg=SURFACE, pady=16)
        hdr.pack(fill="x")
        tk.Label(hdr, text="Smart OCR Tool",
                 font=("Segoe UI", 22, "bold"),
                 fg=ACCENT, bg=SURFACE).pack()
        tk.Label(hdr,
                 text="Image  ·  PDF  ·  DOCX  ·  TXT  —  Extract text & Copy, Save or Listen",
                 font=F_SMALL, fg=MUTED, bg=SURFACE).pack(pady=(3, 0))

    # ── Scrollable body ───────────────────────────────────────────────
    def _setup_scroll_area(self):
        self._canvas = tk.Canvas(self.root, bg=BG, highlightthickness=0)
        sb = tk.Scrollbar(self.root, orient="vertical",
                          command=self._canvas.yview)
        self._canvas.configure(yscrollcommand=sb.set)
        sb.pack(side="right", fill="y")
        self._canvas.pack(side="left", fill="both", expand=True)

        self._inner = tk.Frame(self._canvas, bg=BG)
        self._win_id = self._canvas.create_window(
            (0, 0), window=self._inner, anchor="nw")

        self._inner.bind("<Configure>",
            lambda e: self._canvas.configure(
                scrollregion=self._canvas.bbox("all")))
        self._canvas.bind("<Configure>",
            lambda e: self._canvas.itemconfig(self._win_id, width=e.width))
        self._canvas.bind_all("<MouseWheel>",
            lambda e: self._canvas.yview_scroll(
                -1 * (e.delta // 120), "units"))

    # ── Upload zone ───────────────────────────────────────────────────
    def _build_upload(self):
        pad = tk.Frame(self._inner, bg=BG, padx=28, pady=18)
        pad.pack(fill="x")

        box = tk.Frame(pad, bg=SUR2, cursor="hand2", pady=28)
        box.pack(fill="x")

        tk.Label(box, text="📂", font=("Segoe UI", 30),
                 bg=SUR2, fg=ACCENT).pack()
        tk.Label(box, text="Click here to browse a file",
                 font=F_HEAD, fg=TEXT, bg=SUR2).pack(pady=(7, 3))
        tk.Label(box,
                 text="PNG  ·  JPG  ·  BMP  ·  TIFF  ·  WEBP  ·  PDF  ·  DOCX  ·  TXT",
                 font=F_SMALL, fg=MUTED, bg=SUR2).pack()

        for w in [box] + list(box.winfo_children()):
            w.bind("<Button-1>", lambda e: self._browse())

        self._fbar = tk.Frame(pad, bg=SUR2)
        self._fname_lbl = tk.Label(
            self._fbar, text="", font=F_BODY,
            fg=TEXT, bg=SUR2, anchor="w", padx=10, pady=8)
        self._fname_lbl.pack(side="left", fill="x", expand=True)
        tk.Button(self._fbar, text="✕ Remove", font=F_SMALL,
                  fg=DANGER, bg=SUR2, bd=0, cursor="hand2",
                  activeforeground=DANGER, activebackground=SUR2,
                  command=self._clear_file).pack(side="right", padx=10)

        self._ext_btn = tk.Button(
            pad, text="⚡  Extract Text",
            font=F_BTN, fg=WHITE, bg=ACCENT,
            activeforeground=WHITE, activebackground="#5a52e0",
            bd=0, pady=12, cursor="hand2",
            state="disabled", command=self._start_extraction)
        self._ext_btn.pack(fill="x", pady=(10, 0))

    # ── Progress bar ──────────────────────────────────────────────────
    def _build_progress(self):
        self._prog_frame = tk.Frame(self._inner, bg=BG, padx=28)
        self._prog_bar = ttk.Progressbar(
            self._prog_frame, mode="determinate", style="TProgressbar")
        self._prog_bar.pack(fill="x")
        self._prog_lbl = tk.Label(self._prog_frame, text="",
                                  font=F_SMALL, fg=MUTED, bg=BG)
        self._prog_lbl.pack(pady=(4, 0))

    # ── Output section ────────────────────────────────────────────────
    def _build_output(self):
        self._out_frame = tk.Frame(self._inner, bg=BG, padx=28)

        top = tk.Frame(self._out_frame, bg=BG)
        top.pack(fill="x", pady=(0, 6))
        tk.Label(top, text="EXTRACTED TEXT",
                 font=("Segoe UI", 8, "bold"), fg=MUTED, bg=BG).pack(side="left")
        self._wc_lbl = tk.Label(top, text="", font=F_SMALL, fg=MUTED, bg=BG)
        self._wc_lbl.pack(side="right")

        wrap = tk.Frame(self._out_frame, bg=SURFACE, bd=1, relief="flat")
        wrap.pack(fill="both", expand=True)

        self._text_area = tk.Text(
            wrap, font=F_MONO, fg=TEXT, bg=SURFACE,
            insertbackground=TEXT, selectbackground=ACCENT,
            wrap="word", bd=0, padx=12, pady=10,
            height=12, state="disabled")
        self._text_area.pack(side="left", fill="both", expand=True)

        sb = tk.Scrollbar(wrap, command=self._text_area.yview,
                          bg=SUR2, troughcolor=SUR2)
        sb.pack(side="right", fill="y")
        self._text_area.config(yscrollcommand=sb.set)

        # ── Action buttons ────────────────────────────────────────────
        self._btn_row = tk.Frame(self._out_frame, bg=BG, pady=10)
        self._btn_row.pack(fill="x")

        self._copy_btn = self._mk_btn(
            self._btn_row, "📋  Copy", self._copy, fg=TEXT, bg=SUR2)
        self._copy_btn.pack(side="left", padx=(0, 8))

        self._save_btn = self._mk_btn(
            self._btn_row, "💾  Save", self._save, fg=TEXT, bg=SUR2)
        self._save_btn.pack(side="left", padx=(0, 8))

        # Read — visible initially
        self._read_btn = self._mk_btn(
            self._btn_row, "🔊  Read", self._on_read, fg=WHITE, bg=ACCENT)
        self._read_btn.pack(side="left", padx=(0, 8))

        # Stop — hidden until reading starts
        self._stop_btn = self._mk_btn(
            self._btn_row, "⏹  Stop", self._on_stop, fg=WHITE, bg=DANGER)

        # Resume — hidden until stopped
        self._resume_btn = self._mk_btn(
            self._btn_row, "▶  Resume", self._on_resume, fg=WHITE, bg="#2d7a4f")

        # Status label
        self._tts_lbl = tk.Label(
            self._btn_row, text="", font=F_SMALL, fg=ACCENT2, bg=BG)
        self._tts_lbl.pack(side="left", padx=(10, 0))

    # ── Button factory ────────────────────────────────────────────────
    def _mk_btn(self, parent, text, cmd, fg=TEXT, bg=SUR2):
        return tk.Button(
            parent, text=text, font=F_SMALL,
            fg=fg, bg=bg,
            activeforeground=fg, activebackground=bg,
            bd=0, relief="flat", padx=14, pady=8,
            cursor="hand2", command=cmd)

    # ── File handling ─────────────────────────────────────────────────
    def _browse(self):
        path = filedialog.askopenfilename(
            title="Select a file",
            filetypes=[
                ("All supported",
                 "*.png *.jpg *.jpeg *.bmp *.tiff *.tif *.webp "
                 "*.pdf *.docx *.doc *.txt"),
                ("Images",        "*.png *.jpg *.jpeg *.bmp *.tiff *.tif *.webp"),
                ("PDF",           "*.pdf"),
                ("Word Document", "*.docx *.doc"),
                ("Text File",     "*.txt"),
                ("All Files",     "*.*"),
            ])
        if path:
            self._set_file(path)

    def _set_file(self, path):
        self.filepath = path
        name = os.path.basename(path)
        kb   = os.path.getsize(path) / 1024
        size = f"{kb / 1024:.1f} MB" if kb > 1024 else f"{kb:.0f} KB"
        self._fname_lbl.config(text=f"📄  {name}   ({size})")
        self._fbar.pack(fill="x", pady=(8, 0))
        self._ext_btn.config(state="normal")
        self._hide_output()

    def _clear_file(self):
        self.filepath = None
        self._fbar.pack_forget()
        self._ext_btn.config(state="disabled")
        self._hide_output()
        self.tts.reset()

    # ── Extraction ────────────────────────────────────────────────────
    def _start_extraction(self):
        if not self.filepath:
            return
        self._ext_btn.config(state="disabled")
        self._hide_output()
        self._prog_frame.pack(fill="x", pady=(0, 8))
        self._prog_bar["value"] = 0

        def run():
            try:
                text = extract_text(self.filepath,
                                    progress_cb=self._update_progress)
                self.root.after(0, self._show_result, text)
            except Exception as e:
                self.root.after(0, self._show_error, str(e))

        threading.Thread(target=run, daemon=True).start()

    def _update_progress(self, pct, msg):
        self.root.after(0, lambda: self._prog_bar.configure(value=pct))
        self.root.after(0, lambda: self._prog_lbl.config(text=msg))

    def _show_result(self, text):
        self._prog_frame.pack_forget()
        self._ext_btn.config(state="normal")

        if not text.strip():
            messagebox.showwarning("No Text Found",
                "No text could be extracted.\n\n"
                "Tips:\n"
                "  • Use a higher resolution image\n"
                "  • Ensure Tesseract OCR is installed\n"
                "  • Avoid very blurry images")
            return

        self._text_area.config(state="normal")
        self._text_area.delete("1.0", "end")
        self._text_area.insert("1.0", text)
        self._text_area.config(state="disabled")

        words = len(text.split())
        self._wc_lbl.config(text=f"{words} words  ·  {len(text)} characters")

        self.tts.reset()
        self._set_state_ready()
        self._out_frame.pack(fill="both", expand=True, pady=(10, 0))

    def _show_error(self, msg):
        self._prog_frame.pack_forget()
        self._ext_btn.config(state="normal")
        messagebox.showerror("Extraction Failed",
                             f"{msg}\n\n"
                             "Checklist:\n"
                             "  1. Tesseract installed from:\n"
                             "     github.com/UB-Mannheim/tesseract/wiki\n"
                             "  2. Path in code matches your install\n"
                             "  3. All pip libraries installed in Thonny")

    def _hide_output(self):
        self._out_frame.pack_forget()
        self.tts.reset()

    # ── Copy & Save ───────────────────────────────────────────────────
    def _get_text(self):
        return self._text_area.get("1.0", "end").strip()

    def _copy(self):
        t = self._get_text()
        if t:
            self.root.clipboard_clear()
            self.root.clipboard_append(t)
            messagebox.showinfo("Copied!", "Text copied to clipboard.")

    def _save(self):
        t = self._get_text()
        if not t:
            return
        base = os.path.splitext(
            os.path.basename(self.filepath or "extracted"))[0]
        path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            initialfile=f"{base}_extracted.txt",
            filetypes=[("Text File", "*.txt"), ("All Files", "*.*")])
        if path:
            with open(path, "w", encoding="utf-8") as f:
                f.write(t)
            messagebox.showinfo("Saved!", f"File saved to:\n{path}")

    # ── TTS button states ─────────────────────────────────────────────
    def _set_state_ready(self):
        """Show: Read. Hide: Stop, Resume. Label: empty."""
        self._stop_btn.pack_forget()
        self._resume_btn.pack_forget()
        self._read_btn.pack(side="left", padx=(0, 8))
        self._tts_lbl.config(text="")

    def _set_state_reading(self):
        """Show: Stop. Hide: Read, Resume. Label: Reading..."""
        self._read_btn.pack_forget()
        self._resume_btn.pack_forget()
        self._stop_btn.pack(side="left", padx=(0, 8))
        self._tts_lbl.config(text="Reading...", fg=ACCENT2)

    def _set_state_stopped(self, sentence_num, total):
        """Show: Resume. Hide: Read, Stop. Label: Stopped at sentence N."""
        self._read_btn.pack_forget()
        self._stop_btn.pack_forget()
        self._resume_btn.pack(side="left", padx=(0, 8))
        self._tts_lbl.config(
            text=f"Stopped  (sentence {sentence_num} of {total})",
            fg=WARNING)

    # ── TTS actions ───────────────────────────────────────────────────
    def _on_read(self):
        t = self._get_text()
        if not t:
            return
        self.tts.reset()          # reset position to 0 for fresh read
        self.tts.speak(
            t,
            on_done=lambda: self.root.after(0, self._on_tts_finished)
        )
        self._set_state_reading()

    def _on_stop(self):
        self.tts.stop()
        # Show which sentence we stopped at
        idx   = self.tts._current_index + 1          # 1-based for display
        total = len(self.tts._sentences)
        self._set_state_stopped(idx, total)

    def _on_resume(self):
        self.tts.resume(
            on_done=lambda: self.root.after(0, self._on_tts_finished)
        )
        self._set_state_reading()

    def _on_tts_finished(self):
        self._set_state_ready()
        self._tts_lbl.config(text="Done ✓", fg=ACCENT2)
        self.root.after(2000, lambda: self._tts_lbl.config(text=""))


# ════════════════════════════════════════════════════════════════════
#  ENTRY POINT
# ════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    root = tk.Tk()
    App(root)
    root.mainloop()